"""Генерация DOCX-отчёта через python-docx."""

import io
from typing import Any

from docx import Document


def generate_docx(context: dict[str, Any]) -> bytes:
    doc = Document()
    project = context["project"]
    enabled = set(
        context.get("sections") or ["summary", "pipes", "tanks", "electrical", "specification"]
    )
    doc.add_heading(f"HeatCalc — {project['name']}", level=1)
    if project.get("description"):
        doc.add_paragraph(project["description"])

    objects_section = enabled & {"pipes", "tanks", "electrical", "summary"}
    if objects_section:
        doc.add_heading("Объекты", level=2)
    objects = context.get("objects", []) if objects_section else []
    if objects:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "Тип"
        hdr[2].text = "Параметры"
        hdr[3].text = "Результат"
        for idx, o in enumerate(objects, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = o["object_type"]
            row[2].text = str(o.get("params", {}))
            row[3].text = str(o.get("results") or "—")
    elif objects_section:
        doc.add_paragraph("Объектов нет")

    if "specification" in enabled:
        doc.add_heading("Спецификация", level=2)
    items = context.get("specification", {}).get("items", []) if "specification" in enabled else []
    if items:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Категория"
        hdr[1].text = "Наименование"
        hdr[2].text = "Артикул"
        hdr[3].text = "Ед."
        hdr[4].text = "Кол-во"
        for item in items:
            row = table.add_row().cells
            row[0].text = str(item.get("category", ""))
            row[1].text = str(item.get("name", ""))
            row[2].text = str(item.get("article") or "")
            row[3].text = str(item.get("unit", ""))
            row[4].text = str(item.get("quantity", ""))
    elif "specification" in enabled:
        doc.add_paragraph("Спецификация не сформирована")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
